from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


class CVGeneratorService:
    """Servicio para generar CVs en PDF o Word con todas las secciones del docente"""

    DEFAULT_SECTIONS = [
        'datos_generales',
        'experiencia_laboral',
        'formacion_academica',
        'articulos',
        'libros',
        'congresos',
        'cursos',
        'proyectos',
        'tesis',
        'desarrollos'
    ]

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    @staticmethod
    def _include(section, selected):
        return section in (selected or CVGeneratorService.DEFAULT_SECTIONS)

    @staticmethod
    def _format_date(value):
        return value.strftime('%Y-%m-%d') if value else 'Sin fecha'

    @classmethod
    def _format_period(cls, inicio, fin):
        start = cls._format_date(inicio)
        end = fin.strftime('%Y-%m-%d') if fin else 'Actualidad'
        return f"{start} - {end}"

    @staticmethod
    def _docente_contact(docente):
        correo = docente.correo_principal or 'Correo no registrado'
        nacionalidad = docente.nacionalidad or 'Nacionalidad no registrada'
        return correo, nacionalidad

    # ------------------------------------------------------------------
    # Generación PDF
    # ------------------------------------------------------------------
    def generar_pdf(
        self,
        docente,
        formaciones=None,
        empleos=None,
        articulos=None,
        libros=None,
        congresos=None,
        cursos=None,
        proyectos=None,
        tesis=None,
        desarrollos=None,
        tipo_cv='academico',
        sections=None
    ):
        """Genera un CV en PDF con todas las secciones solicitadas"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2563eb'),
            spaceAfter=6,
            spaceBefore=12
        )

        correo, nacionalidad = self._docente_contact(docente)
        story.append(Paragraph(docente.nombre_completo or 'Docente', title_style))
        story.append(Paragraph(correo, styles['Normal']))
        if docente.orcid:
            story.append(Paragraph(f"ORCID: {docente.orcid}", styles['Normal']))
        story.append(Paragraph(f"Nacionalidad: {nacionalidad}", styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))

        # Datos generales
        if self._include('datos_generales', sections):
            story.append(Paragraph('DATOS GENERALES', heading_style))
            datos = [
                f"<b>CVU:</b> {docente.cvu or 'Sin capturar'}",
                f"<b>RFC:</b> {docente.rfc or 'Sin capturar'}",
                f"<b>CURP:</b> {docente.curp or 'Sin capturar'}",
                f"<b>Researcher ID:</b> {docente.researcher_id or 'Sin capturar'}"
            ]
            for dato in datos:
                story.append(Paragraph(dato, styles['Normal']))
            story.append(Spacer(1, 0.2 * inch))

        # Formación académica
        if self._include('formacion_academica', sections) and formaciones:
            story.append(Paragraph('FORMACIÓN ACADÉMICA', heading_style))
            for grado in formaciones:
                titulo = grado.grado_obtenido or grado.nivel
                cuerpo = f"<b>{titulo}</b> - {grado.institucion or 'Institución no registrada'}"
                if grado.pais:
                    cuerpo += f" ({grado.pais})"
                cuerpo += f"<br/>{self._format_period(grado.fecha_inicio, grado.fecha_fin)}"
                if grado.area_conocimiento:
                    cuerpo += f"<br/>Área: {grado.area_conocimiento}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Experiencia laboral
        if self._include('experiencia_laboral', sections) and empleos:
            story.append(Paragraph('EXPERIENCIA LABORAL', heading_style))
            for empleo in empleos:
                cuerpo = f"<b>{empleo.puesto}</b> · {empleo.institucion}"
                cuerpo += f"<br/>{self._format_period(empleo.fecha_inicio, empleo.fecha_fin)}"
                if empleo.logros:
                    cuerpo += f"<br/>Logros: {empleo.logros}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Artículos
        if self._include('articulos', sections) and articulos:
            story.append(Paragraph('ARTÍCULOS CIENTÍFICOS', heading_style))
            for idx, articulo in enumerate(articulos, 1):
                cuerpo = f"<b>{idx}. {articulo.titulo}</b>"
                if articulo.revista:
                    cuerpo += f"<br/>{articulo.revista}"
                if articulo.anio:
                    cuerpo += f", {articulo.anio}"
                if articulo.doi:
                    cuerpo += f"<br/>DOI: {articulo.doi}"
                if articulo.autores:
                    cuerpo += f"<br/>Autores: {articulo.autores}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Libros
        if self._include('libros', sections) and libros:
            story.append(Paragraph('LIBROS Y CAPÍTULOS', heading_style))
            for libro in libros:
                cuerpo = f"<b>{libro.titulo}</b>"
                if libro.titulo_capitulo:
                    cuerpo += f" · Capítulo: {libro.titulo_capitulo}"
                cuerpo += f"<br/>{libro.editorial or 'Editorial no registrada'}"
                if libro.anio:
                    cuerpo += f", {libro.anio}"
                if libro.isbn:
                    cuerpo += f"<br/>ISBN: {libro.isbn}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Congresos
        if self._include('congresos', sections) and congresos:
            story.append(Paragraph('CONGRESOS Y PONENCIAS', heading_style))
            for congreso in congresos:
                cuerpo = f"<b>{congreso.nombre_congreso}</b>"
                if congreso.titulo_ponencia:
                    cuerpo += f"<br/>Ponencia: {congreso.titulo_ponencia}"
                if congreso.fecha:
                    cuerpo += f"<br/>{self._format_date(congreso.fecha)}"
                ubicacion = ', '.join(filter(None, [congreso.ciudad, congreso.pais]))
                if ubicacion:
                    cuerpo += f" · {ubicacion}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Cursos impartidos
        if self._include('cursos', sections) and cursos:
            story.append(Paragraph('CURSOS IMPARTIDOS', heading_style))
            for curso in cursos:
                cuerpo = f"<b>{curso.nombre_curso}</b>"
                if curso.programa_educativo:
                    cuerpo += f" · {curso.programa_educativo}"
                cuerpo += f"<br/>{self._format_period(curso.fecha_inicio, curso.fecha_fin)}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Proyectos
        if self._include('proyectos', sections) and proyectos:
            story.append(Paragraph('PROYECTOS DE INVESTIGACIÓN', heading_style))
            for proyecto in proyectos:
                cuerpo = f"<b>{proyecto.nombre_proyecto}</b>"
                if proyecto.linea_investigacion:
                    cuerpo += f" · Línea: {proyecto.linea_investigacion}"
                cuerpo += f"<br/>{self._format_period(proyecto.fecha_inicio, proyecto.fecha_fin)}"
                if proyecto.estado:
                    cuerpo += f" · Estado: {proyecto.estado}"
                if proyecto.objetivo_general:
                    cuerpo += f"<br/>Objetivo: {proyecto.objetivo_general}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Tesis
        if self._include('tesis', sections) and tesis:
            story.append(Paragraph('TESIS DIRIGIDAS', heading_style))
            for trabajo in tesis:
                cuerpo = f"<b>{trabajo.titulo}</b>"
                if trabajo.estudiante_nombre:
                    cuerpo += f" · Estudiante: {trabajo.estudiante_nombre}"
                if trabajo.nivel:
                    cuerpo += f"<br/>Nivel: {trabajo.nivel}"
                cuerpo += f"<br/>{trabajo.institucion or 'Institución no registrada'}"
                cuerpo += f" · {self._format_period(trabajo.fecha_inicio, trabajo.fecha_fin)}"
                if trabajo.estado:
                    cuerpo += f" · Estado: {trabajo.estado}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        # Desarrollos tecnológicos
        if self._include('desarrollos', sections) and desarrollos:
            story.append(Paragraph('DESARROLLOS TECNOLÓGICOS', heading_style))
            for desarrollo in desarrollos:
                cuerpo = f"<b>{desarrollo.nombre}</b>"
                if desarrollo.tipo:
                    cuerpo += f" · Tipo: {desarrollo.tipo}"
                if desarrollo.nivel_madurez:
                    cuerpo += f"<br/>Nivel de madurez: {desarrollo.nivel_madurez}"
                if desarrollo.descripcion:
                    cuerpo += f"<br/>{desarrollo.descripcion}"
                story.append(Paragraph(cuerpo, styles['Normal']))
                story.append(Spacer(1, 0.08 * inch))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    # ------------------------------------------------------------------
    # Generación Word
    # ------------------------------------------------------------------
    def generar_word(
        self,
        docente,
        formaciones=None,
        empleos=None,
        articulos=None,
        libros=None,
        congresos=None,
        cursos=None,
        proyectos=None,
        tesis=None,
        desarrollos=None,
        tipo_cv='academico',
        sections=None
    ):
        """Genera un CV en formato Word (.docx)"""
        doc = Document()
        correo, nacionalidad = self._docente_contact(docente)

        heading = doc.add_heading(docente.nombre_completo or 'Docente', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(correo)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Nacionalidad: {nacionalidad}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        if docente.orcid:
            doc.add_paragraph(f"ORCID: {docente.orcid}").alignment = WD_ALIGN_PARAGRAPH.CENTER

        def add_section(title):
            doc.add_heading(title, level=1)

        if self._include('datos_generales', sections):
            add_section('Datos Generales')
            doc.add_paragraph(f"CVU: {docente.cvu or 'Sin capturar'}")
            doc.add_paragraph(f"RFC: {docente.rfc or 'Sin capturar'}")
            doc.add_paragraph(f"CURP: {docente.curp or 'Sin capturar'}")
            doc.add_paragraph(f"Researcher ID: {docente.researcher_id or 'Sin capturar'}")

        if self._include('formacion_academica', sections) and formaciones:
            add_section('Formación Académica')
            for grado in formaciones:
                titulo = grado.grado_obtenido or grado.nivel
                doc.add_paragraph(f"{titulo} - {grado.institucion or 'Institución no registrada'}")
                doc.add_paragraph(f"Periodo: {self._format_period(grado.fecha_inicio, grado.fecha_fin)}")
                if grado.area_conocimiento:
                    doc.add_paragraph(f"Área: {grado.area_conocimiento}")

        if self._include('experiencia_laboral', sections) and empleos:
            add_section('Experiencia Laboral')
            for empleo in empleos:
                doc.add_paragraph(f"{empleo.puesto} · {empleo.institucion}")
                doc.add_paragraph(f"Periodo: {self._format_period(empleo.fecha_inicio, empleo.fecha_fin)}")
                if empleo.logros:
                    doc.add_paragraph(f"Logros: {empleo.logros}")

        if self._include('articulos', sections) and articulos:
            add_section('Artículos Científicos')
            for articulo in articulos:
                doc.add_paragraph(articulo.titulo, style='List Number')
                if articulo.revista:
                    doc.add_paragraph(f"Revista: {articulo.revista}")
                if articulo.anio:
                    doc.add_paragraph(f"Año: {articulo.anio}")
                if articulo.doi:
                    doc.add_paragraph(f"DOI: {articulo.doi}")
                if articulo.autores:
                    doc.add_paragraph(f"Autores: {articulo.autores}")

        if self._include('libros', sections) and libros:
            add_section('Libros y Capítulos')
            for libro in libros:
                doc.add_paragraph(libro.titulo, style='List Number')
                if libro.titulo_capitulo:
                    doc.add_paragraph(f"Capítulo: {libro.titulo_capitulo}")
                if libro.editorial:
                    doc.add_paragraph(f"Editorial: {libro.editorial}")
                if libro.anio:
                    doc.add_paragraph(f"Año: {libro.anio}")
                if libro.isbn:
                    doc.add_paragraph(f"ISBN: {libro.isbn}")

        if self._include('congresos', sections) and congresos:
            add_section('Congresos y Ponencias')
            for congreso in congresos:
                doc.add_paragraph(congreso.nombre_congreso, style='List Bullet')
                if congreso.titulo_ponencia:
                    doc.add_paragraph(f"Ponencia: {congreso.titulo_ponencia}")
                if congreso.fecha:
                    doc.add_paragraph(f"Fecha: {self._format_date(congreso.fecha)}")
                ubicacion = ', '.join(filter(None, [congreso.ciudad, congreso.pais]))
                if ubicacion:
                    doc.add_paragraph(f"Ubicación: {ubicacion}")

        if self._include('cursos', sections) and cursos:
            add_section('Cursos Impartidos')
            for curso in cursos:
                doc.add_paragraph(curso.nombre_curso, style='List Bullet')
                if curso.programa_educativo:
                    doc.add_paragraph(f"Programa: {curso.programa_educativo}")
                doc.add_paragraph(f"Periodo: {self._format_period(curso.fecha_inicio, curso.fecha_fin)}")

        if self._include('proyectos', sections) and proyectos:
            add_section('Proyectos de Investigación')
            for proyecto in proyectos:
                doc.add_paragraph(proyecto.nombre_proyecto, style='List Number')
                if proyecto.linea_investigacion:
                    doc.add_paragraph(f"Línea: {proyecto.linea_investigacion}")
                if proyecto.estado:
                    doc.add_paragraph(f"Estado: {proyecto.estado}")
                doc.add_paragraph(f"Periodo: {self._format_period(proyecto.fecha_inicio, proyecto.fecha_fin)}")
                if proyecto.objetivo_general:
                    doc.add_paragraph(f"Objetivo: {proyecto.objetivo_general}")

        if self._include('tesis', sections) and tesis:
            add_section('Tesis Dirigidas')
            for trabajo in tesis:
                doc.add_paragraph(trabajo.titulo, style='List Bullet')
                if trabajo.estudiante_nombre:
                    doc.add_paragraph(f"Estudiante: {trabajo.estudiante_nombre}")
                doc.add_paragraph(f"Institución: {trabajo.institucion or 'No registrada'}")
                doc.add_paragraph(f"Periodo: {self._format_period(trabajo.fecha_inicio, trabajo.fecha_fin)}")
                if trabajo.estado:
                    doc.add_paragraph(f"Estado: {trabajo.estado}")

        if self._include('desarrollos', sections) and desarrollos:
            add_section('Desarrollos Tecnológicos')
            for desarrollo in desarrollos:
                doc.add_paragraph(desarrollo.nombre, style='List Bullet')
                if desarrollo.tipo:
                    doc.add_paragraph(f"Tipo: {desarrollo.tipo}")
                if desarrollo.nivel_madurez:
                    doc.add_paragraph(f"Nivel de madurez: {desarrollo.nivel_madurez}")
                if desarrollo.descripcion:
                    doc.add_paragraph(desarrollo.descripcion)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
